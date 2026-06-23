"""
TrueVow DRAFT™ Service - Document Parser
Parses PDF and DOCX documents for validation
"""

import io
import logging
from typing import Dict, Any, Optional
import PyPDF2
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


def parse_document(
    content: bytes,
    mime_type: str,
    filename: Optional[str] = None
) -> Dict[str, Any]:
    """
    Parse document content and extract text for validation.
    
    Args:
        content: Document content as bytes
        mime_type: MIME type of the document
        filename: Optional filename for better error messages
    
    Returns:
        Dictionary with:
        - parse_success: bool
        - text: str (extracted text)
        - pages: List[str] (for PDFs)
        - metadata: Dict with page_count, word_count, etc.
        - error: Optional[str] (if parsing failed)
    """
    try:
        if mime_type == "application/pdf" or (filename and filename.endswith(".pdf")):
            return _parse_pdf(content, filename)
        elif mime_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"
        ] or (filename and (filename.endswith(".docx") or filename.endswith(".doc"))):
            return _parse_docx(content, filename)
        elif mime_type == "text/plain" or (filename and filename.endswith(".txt")):
            return _parse_text(content, filename)
        else:
            return {
                "parse_success": False,
                "error": f"Unsupported file type: {mime_type}",
                "text": "",
                "metadata": {}
            }
    except Exception as e:
        logger.error(f"Error parsing document: {e}", exc_info=True)
        return {
            "parse_success": False,
            "error": str(e),
            "text": "",
            "metadata": {}
        }


def _parse_pdf(content: bytes, filename: Optional[str] = None) -> Dict[str, Any]:
    """Parse PDF document"""
    try:
        pdf_file = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        pages = []
        full_text = ""
        
        for page_num, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            pages.append(page_text)
            full_text += page_text + "\n"
        
        word_count = len(full_text.split())
        character_count = len(full_text)
        
        return {
            "parse_success": True,
            "text": full_text.strip(),
            "pages": pages,
            "metadata": {
                "page_count": len(pdf_reader.pages),
                "word_count": word_count,
                "character_count": character_count,
                "file_type": "application/pdf"
            }
        }
    except Exception as e:
        logger.error(f"Error parsing PDF: {e}", exc_info=True)
        return {
            "parse_success": False,
            "error": f"Failed to parse PDF: {str(e)}",
            "text": "",
            "metadata": {}
        }


def _parse_docx(content: bytes, filename: Optional[str] = None) -> Dict[str, Any]:
    """Parse DOCX document"""
    try:
        # DOCX files are ZIP archives, check if it's a valid ZIP first
        import zipfile
        try:
            zipfile.ZipFile(io.BytesIO(content))
        except zipfile.BadZipFile:
            return {
                "parse_success": False,
                "error": "Invalid DOCX file: Not a valid ZIP archive",
                "text": "",
                "metadata": {}
            }
        
        docx_file = io.BytesIO(content)
        doc = DocxDocument(docx_file)
        
        paragraphs = []
        full_text = ""
        
        # Extract paragraphs
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                paragraphs.append(para_text)
                full_text += para_text + "\n"
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    full_text += row_text + "\n"
        
        word_count = len(full_text.split())
        character_count = len(full_text)
        
        # Check for header/footer
        has_header = len(doc.sections) > 0 and doc.sections[0].header is not None
        
        return {
            "parse_success": True,
            "text": full_text.strip(),
            "paragraphs": paragraphs,
            "metadata": {
                "paragraph_count": len(paragraphs),
                "word_count": word_count,
                "character_count": character_count,
                "has_header": has_header,
                "file_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            }
        }
    except Exception as e:
        logger.error(f"Error parsing DOCX: {e}", exc_info=True)
        return {
            "parse_success": False,
            "error": f"Failed to parse DOCX: {str(e)}",
            "text": "",
            "metadata": {}
        }


def _parse_text(content: bytes, filename: Optional[str] = None) -> Dict[str, Any]:
    """Parse plain text document"""
    try:
        # Try UTF-8 first, then Latin-1
        try:
            text = content.decode('utf-8')
        except UnicodeDecodeError:
            text = content.decode('latin-1')
        
        word_count = len(text.split())
        character_count = len(text)
        
        return {
            "parse_success": True,
            "text": text,
            "metadata": {
                "word_count": word_count,
                "character_count": character_count,
                "file_type": "text/plain"
            }
        }
    except Exception as e:
        logger.error(f"Error parsing text: {e}", exc_info=True)
        return {
            "parse_success": False,
            "error": f"Failed to parse text: {str(e)}",
            "text": "",
            "metadata": {}
        }

